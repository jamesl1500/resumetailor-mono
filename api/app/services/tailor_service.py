import html
import io
import os
import re
import textwrap
from collections import Counter
from typing import Optional
from uuid import UUID, uuid4
from httpx import get
from pdfminer.high_level import extract_text as extract_pdf_text
from reportlab.lib import colors
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from sqlalchemy.orm import Session
from app.config.config import Config
from app.models.job_analysis_model import JobAnalysis
from app.models.resume_profile_model import ResumeProfile
from app.models.tailored_resume_model import TailoredResume
from app.services.ai_service import analyze_job_with_openai, parse_resume_with_openai
from docx.shared import Pt, Inches, RGBColor

STOPWORDS = {
    "the",
    "and",
    "for",
    "with",
    "that",
    "this",
    "from",
    "your",
    "you",
    "our",
    "are",
    "will",
    "their",
    "about",
    "into",
    "role",
    "team",
    "work",
    "workflows",
    "ability",
    "strong",
    "experience",
    "years",
    "year",
    "all",
    "have",
    "has",
    "job",
    "across",
    "including",
}

ROLE_SIGNALS = [
    "junior",
    "mid",
    "senior",
    "lead",
    "principal",
    "manager",
    "director",
    "head",
    "vp",
]

TOOL_SIGNALS = [
    "figma",
    "sketch",
    "adobe",
    "notion",
    "jira",
    "sql",
    "python",
    "react",
    "typescript",
    "node",
]

FOCUS_SIGNALS = [
    "accessibility",
    "experimentation",
    "research",
    "strategy",
    "design systems",
    "stakeholder",
    "metrics",
    "prototype",
    "ui",
    "ux",
]

EMAIL_PATTERN = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_PATTERN = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")

SECTION_HEADERS = {
    "experience",
    "work experience",
    "employment",
    "education",
    "skills",
    "projects",
    "summary",
}

DATE_PATTERN = re.compile(
    r"(\b(?:\d{4}|\w{3,9}\s+\d{4})\b\s*(?:-|to)\s*\b(?:present|current|\d{4}|\w{3,9}\s+\d{4})\b)",
    re.IGNORECASE,
)


def _slice_section(lines: list[str], header: str) -> list[str]:
    header_lower = header.lower()
    start_idx = None
    for idx, line in enumerate(lines):
        if line.lower().strip(" :") == header_lower:
            start_idx = idx + 1
            break

    if start_idx is None:
        return []

    section_lines = []
    for line in lines[start_idx:]:
        if line.lower().strip(" :") in SECTION_HEADERS:
            break
        section_lines.append(line)

    return section_lines


def _extract_bullets(section_lines: list[str]) -> list[str]:
    bullets = []
    for line in section_lines:
        if line.startswith("-"):
            bullets.append(line.lstrip("- "))
    return bullets


def _split_entries(section_lines: list[str]) -> list[dict]:
    entries: list[dict] = []
    current: dict = {"header": "", "bullets": []}

    for line in section_lines:
        if line.startswith("-"):
            current["bullets"].append(line.lstrip("- "))
            continue

        if current["header"]:
            entries.append(current)
            current = {"header": "", "bullets": []}
        current["header"] = line

    if current["header"] or current["bullets"]:
        entries.append(current)

    return entries


def _parse_header_parts(header: str) -> tuple[str, str, str | None, str | None]:
    title = header
    company = ""
    location = None
    date_range = DATE_PATTERN.search(header)
    if date_range:
        header = header.replace(date_range.group(0), "").strip("- ")
        dates = date_range.group(0)
    else:
        dates = None

    if " - " in header:
        title, company = [part.strip() for part in header.split(" - ", 1)]
    elif " | " in header:
        title, company = [part.strip() for part in header.split(" | ", 1)]
    elif "," in header:
        title, company = [part.strip() for part in header.split(",", 1)]

    if company and "," in company:
        company, location = [part.strip() for part in company.split(",", 1)]

    start_date = None
    end_date = None
    if dates:
        parts = re.split(r"\s*(?:-|to)\s*", dates, maxsplit=1, flags=re.IGNORECASE)
        if parts:
            start_date = parts[0].strip()
        if len(parts) > 1:
            end_date = parts[1].strip()

    return title or "Role", company or "Company", start_date, end_date


def _parse_education_header(header: str) -> tuple[str, str | None, str | None]:
    institution = header
    degree = None
    field = None

    date_range = DATE_PATTERN.search(header)
    if date_range:
        header = header.replace(date_range.group(0), "").strip("- ")

    if " - " in header:
        institution, rest = [part.strip() for part in header.split(" - ", 1)]
        if "," in rest:
            degree, field = [part.strip() for part in rest.split(",", 1)]
        else:
            degree = rest
    elif "," in header:
        institution, rest = [part.strip() for part in header.split(",", 1)]
        if "," in rest:
            degree, field = [part.strip() for part in rest.split(",", 1)]
        else:
            degree = rest

    return institution or "University", degree, field


def _extract_experience_section(lines: list[str]) -> list[dict]:
    section_lines = _slice_section(lines, "Experience")
    if not section_lines:
        section_lines = _slice_section(lines, "Work Experience")

    entries = _split_entries(section_lines)
    if not entries:
        return []

    parsed = []
    for entry in entries:
        header = entry.get("header", "")
        bullets = entry.get("bullets", [])
        title, company, start_date, end_date = _parse_header_parts(header)
        location = None
        if "," in header and company and company in header:
            trailing = header.split(company, 1)[-1]
            if "," in trailing:
                location = trailing.split(",", 1)[-1].strip() or None
        parsed.append(
            {
                "title": title,
                "company": company,
                "location": location,
                "start_date": start_date,
                "end_date": end_date,
                "bullets": bullets or ["Led cross-functional initiatives"],
            }
        )

    return parsed


def _extract_education_section(lines: list[str]) -> list[dict]:
    section_lines = _slice_section(lines, "Education")
    if not section_lines:
        return []

    entries = _split_entries(section_lines)
    if not entries:
        return []

    parsed = []
    for entry in entries:
        header = entry.get("header", "")
        bullets = entry.get("bullets", [])
        institution, degree, field = _parse_education_header(header)
        start_date = None
        end_date = None
        date_range = DATE_PATTERN.search(header)
        if date_range:
            parts = re.split(r"\s*(?:-|to)\s*", date_range.group(0), maxsplit=1, flags=re.IGNORECASE)
            if parts:
                start_date = parts[0].strip()
            if len(parts) > 1:
                end_date = parts[1].strip()
        parsed.append(
            {
                "institution": institution,
                "degree": degree,
                "field_of_study": field,
                "start_date": start_date,
                "end_date": end_date,
                "bullets": bullets or ["Relevant coursework and projects"],
            }
        )

    return parsed


def strip_html(text: str) -> str:
    without_tags = re.sub(r"<[^>]+>", " ", text)
    return html.unescape(without_tags)


def normalize_text(text: str) -> str:
    cleaned = strip_html(text)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def extract_keywords(text: str, limit: int = 10) -> list[str]:
    tokens = re.findall(r"[A-Za-z][A-Za-z+#.-]{2,}", text.lower())
    filtered = [token for token in tokens if token not in STOPWORDS]
    counts = Counter(filtered)
    return [word for word, _ in counts.most_common(limit)]


def extract_signals(text: str) -> dict[str, list[str]]:
    lower = text.lower()
    levels = [signal for signal in ROLE_SIGNALS if signal in lower]
    tools = [signal for signal in TOOL_SIGNALS if signal in lower]
    focus = [signal for signal in FOCUS_SIGNALS if signal in lower]
    return {
        "levels": list(dict.fromkeys(levels)),
        "tools": list(dict.fromkeys(tools)),
        "focus": list(dict.fromkeys(focus)),
    }


def summarize_text(text: str, max_sentences: int = 2) -> str:
    sentences = re.split(r"(?<=[.!?])\s+", text)
    summary = " ".join(sentences[:max_sentences]).strip()
    return summary or text[:200]


def fetch_job_text(job_url: str) -> str:
    response = get(
        job_url,
        timeout=10.0,
        headers={"User-Agent": "ResumeTailorBot/1.0"},
    )
    response.raise_for_status()
    return response.text


def extract_text_from_pdf(data: bytes) -> str:
    with io.BytesIO(data) as buffer:
        return extract_pdf_text(buffer)


def extract_text_from_docx(data: bytes) -> str:
    with io.BytesIO(data) as buffer:
        document = Document(buffer)
        return "\n".join(paragraph.text for paragraph in document.paragraphs)


def extract_text_from_upload(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if lower.endswith(".docx"):
        return extract_text_from_docx(data)
    raise ValueError("Unsupported file type. Use PDF or DOCX.")


def ensure_storage_dir(*parts: str) -> str:
    base_dir = Config.STORAGE_PATH
    path = os.path.join(base_dir, *parts)
    os.makedirs(path, exist_ok=True)
    return path


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip("-")
    return cleaned or "Resume"


def _style_profile(style: Optional[str]) -> dict:
    name = (style or "Modern").lower()
    if name == "slim":
        return {
            "title_size": 16,
            "subtitle_size": 11,
            "body_size": 10,
            "accent": RGBColor(140, 96, 64),
            "spacing_after": 4,
            "line_spacing": 1.0,
        }
    if name == "fancy":
        return {
            "title_size": 20,
            "subtitle_size": 12,
            "body_size": 11,
            "accent": RGBColor(201, 95, 60),
            "spacing_after": 8,
            "line_spacing": 1.15,
        }
    return {
        "title_size": 18,
        "subtitle_size": 12,
        "body_size": 11,
        "accent": RGBColor(201, 95, 60),
        "spacing_after": 6,
        "line_spacing": 1.1,
    }


def render_docx(
    path: str,
    name: str,
    role: str,
    summary: str,
    bullets: list[str],
    experiences: list[dict],
    education: list[dict],
    style: Optional[str],
) -> None:
    document = Document()

    profile = _style_profile(style)
    base_style = document.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(profile["body_size"])
    base_style.paragraph_format.space_after = Pt(profile["spacing_after"])
    base_style.paragraph_format.line_spacing = profile["line_spacing"]

    title = document.add_paragraph()
    title_run = title.add_run(f"{name}")
    title_run.font.size = Pt(profile["title_size"])
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(45, 45, 45)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(role)
    subtitle_run.font.size = Pt(profile["subtitle_size"])
    subtitle_run.font.color.rgb = RGBColor(120, 120, 120)
    subtitle.paragraph_format.space_after = Pt(10)

    summary_para = document.add_paragraph(summary)
    summary_para.paragraph_format.space_after = Pt(12)

    highlight_heading = document.add_paragraph("Highlights")
    highlight_heading.runs[0].bold = True
    highlight_heading.runs[0].font.size = Pt(profile["subtitle_size"])
    highlight_heading.runs[0].font.color.rgb = profile["accent"]
    highlight_heading.paragraph_format.space_after = Pt(6)

    for bullet in bullets:
        para = document.add_paragraph(bullet, style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.space_after = Pt(4)

    if experiences:
        experience_heading = document.add_paragraph("Experience")
        experience_heading.runs[0].bold = True
        experience_heading.runs[0].font.size = Pt(profile["subtitle_size"])
        experience_heading.runs[0].font.color.rgb = profile["accent"]
        experience_heading.paragraph_format.space_after = Pt(6)
        for experience in experiences:
            title = experience.get("title", "")
            company = experience.get("company", "")
            location = experience.get("location") or ""
            dates = ""
            if experience.get("start_date") or experience.get("end_date"):
                dates = " - ".join(
                    [
                        part
                        for part in [
                            experience.get("start_date"),
                            experience.get("end_date"),
                        ]
                        if part
                    ]
                )
            header = " - ".join([part for part in [title, company] if part])
            if header:
                header_line = document.add_paragraph()
                header_run = header_line.add_run(header)
                header_run.bold = True
                header_run.font.size = Pt(11)
                meta = " · ".join([part for part in [location, dates] if part])
                if meta:
                    meta_run = header_line.add_run(f"  {meta}")
                    meta_run.font.color.rgb = RGBColor(120, 120, 120)
                header_line.paragraph_format.space_after = Pt(4)
            for bullet in experience.get("bullets", []):
                para = document.add_paragraph(bullet, style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.15)
                para.paragraph_format.space_after = Pt(4)

    if education:
        education_heading = document.add_paragraph("Education")
        education_heading.runs[0].bold = True
        education_heading.runs[0].font.size = Pt(profile["subtitle_size"])
        education_heading.runs[0].font.color.rgb = profile["accent"]
        education_heading.paragraph_format.space_after = Pt(6)
        for edu in education:
            institution = edu.get("institution", "")
            degree = edu.get("degree", "")
            field = edu.get("field_of_study", "")
            dates = ""
            if edu.get("start_date") or edu.get("end_date"):
                dates = " - ".join(
                    [
                        part
                        for part in [edu.get("start_date"), edu.get("end_date")]
                        if part
                    ]
                )
            header = " - ".join([part for part in [institution, degree, field] if part])
            if header:
                header_line = document.add_paragraph()
                header_run = header_line.add_run(header)
                header_run.bold = True
                header_run.font.size = Pt(11)
                if dates:
                    meta_run = header_line.add_run(f"  {dates}")
                    meta_run.font.color.rgb = RGBColor(120, 120, 120)
                header_line.paragraph_format.space_after = Pt(4)
            for bullet in edu.get("bullets", []):
                para = document.add_paragraph(bullet, style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.15)
                para.paragraph_format.space_after = Pt(4)
    document.save(path)


def render_pdf(
    path: str,
    name: str,
    role: str,
    summary: str,
    bullets: list[str],
    experiences: list[dict],
    education: list[dict],
    style: Optional[str],
) -> None:
    canvas_obj = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    profile = _style_profile(style)
    margin = 60 if (style or "").lower() == "slim" else 64
    y = height - margin

    accent = (
        colors.HexColor("#8C6040")
        if (style or "").lower() == "slim"
        else colors.HexColor("#C95F3C")
    )
    muted = colors.HexColor("#6A5F55")

    canvas_obj.setFont("Helvetica-Bold", profile["title_size"] + 2)
    canvas_obj.setFillColor(colors.HexColor("#1D1A15"))
    canvas_obj.drawString(margin, y, name)
    y -= 18
    canvas_obj.setFont("Helvetica", profile["subtitle_size"])
    canvas_obj.setFillColor(muted)
    canvas_obj.drawString(margin, y, role)
    y -= 18

    canvas_obj.setStrokeColor(accent)
    canvas_obj.setLineWidth(1)
    canvas_obj.line(margin, y, width - margin, y)
    y -= 18

    canvas_obj.setFont("Helvetica", profile["body_size"])
    canvas_obj.setFillColor(colors.HexColor("#1D1A15"))
    for line in textwrap.wrap(summary, width=95):
        if y < margin:
            canvas_obj.showPage()
            y = height - margin
            canvas_obj.setFont("Helvetica", 11)
        canvas_obj.drawString(margin, y, line)
        y -= 14 if (style or "").lower() == "slim" else 15

    y -= 10
    canvas_obj.setFont("Helvetica-Bold", profile["subtitle_size"])
    canvas_obj.setFillColor(accent)
    canvas_obj.drawString(margin, y, "Highlights")
    y -= 16
    canvas_obj.setFont("Helvetica", profile["body_size"])
    canvas_obj.setFillColor(colors.HexColor("#1D1A15"))

    for bullet in bullets:
        wrapped = textwrap.wrap(bullet, width=92)
        for index, line in enumerate(wrapped):
            if y < margin:
                canvas_obj.showPage()
                y = height - margin
                canvas_obj.setFont("Helvetica", 11)
            prefix = "- " if index == 0 else "  "
            canvas_obj.drawString(margin, y, f"{prefix}{line}")
            y -= 14 if (style or "").lower() == "slim" else 15

    if experiences:
        y -= 8
        canvas_obj.setFont("Helvetica-Bold", profile["subtitle_size"])
        canvas_obj.setFillColor(accent)
        canvas_obj.drawString(margin, y, "Experience")
        y -= 16
        canvas_obj.setFont("Helvetica", profile["body_size"])
        canvas_obj.setFillColor(colors.HexColor("#1D1A15"))
        for experience in experiences:
            header = " - ".join(
                [
                    part
                    for part in [
                        experience.get("title", ""),
                        experience.get("company", ""),
                    ]
                    if part
                ]
            )
            if header:
                canvas_obj.setFont("Helvetica-Bold", profile["body_size"])
                canvas_obj.drawString(margin, y, header)
                y -= 14
                canvas_obj.setFont("Helvetica", max(profile["body_size"] - 1, 9))
                meta = " · ".join(
                    [
                        part
                        for part in [
                            experience.get("location"),
                            " - ".join(
                                [
                                    part
                                    for part in [
                                        experience.get("start_date"),
                                        experience.get("end_date"),
                                    ]
                                    if part
                                ]
                            ),
                        ]
                        if part
                    ]
                )
                if meta:
                    canvas_obj.setFillColor(muted)
                    canvas_obj.drawString(margin, y, meta)
                    canvas_obj.setFillColor(colors.HexColor("#1D1A15"))
                    y -= 12
                canvas_obj.setFont("Helvetica", profile["body_size"])
            for bullet in experience.get("bullets", []):
                wrapped = textwrap.wrap(bullet, width=92)
                for index, line in enumerate(wrapped):
                    if y < margin:
                        canvas_obj.showPage()
                        y = height - margin
                        canvas_obj.setFont("Helvetica", 11)
                    prefix = "- " if index == 0 else "  "
                    canvas_obj.drawString(margin, y, f"{prefix}{line}")
                    y -= 14 if (style or "").lower() == "slim" else 15

    if education:
        y -= 8
        canvas_obj.setFont("Helvetica-Bold", profile["subtitle_size"])
        canvas_obj.setFillColor(accent)
        canvas_obj.drawString(margin, y, "Education")
        y -= 16
        canvas_obj.setFont("Helvetica", profile["body_size"])
        canvas_obj.setFillColor(colors.HexColor("#1D1A15"))
        for edu in education:
            header = " - ".join(
                [
                    part
                    for part in [
                        edu.get("institution", ""),
                        edu.get("degree", ""),
                        edu.get("field_of_study", ""),
                    ]
                    if part
                ]
            )
            if header:
                canvas_obj.setFont("Helvetica-Bold", profile["body_size"])
                canvas_obj.drawString(margin, y, header)
                y -= 14
                dates = " - ".join(
                    [
                        part
                        for part in [edu.get("start_date"), edu.get("end_date")]
                        if part
                    ]
                )
                if dates:
                    canvas_obj.setFont("Helvetica", max(profile["body_size"] - 1, 9))
                    canvas_obj.setFillColor(muted)
                    canvas_obj.drawString(margin, y, dates)
                    canvas_obj.setFillColor(colors.HexColor("#1D1A15"))
                    y -= 12
                canvas_obj.setFont("Helvetica", profile["body_size"])
            for bullet in edu.get("bullets", []):
                wrapped = textwrap.wrap(bullet, width=92)
                for index, line in enumerate(wrapped):
                    if y < margin:
                        canvas_obj.showPage()
                        y = height - margin
                        canvas_obj.setFont("Helvetica", 11)
                    prefix = "- " if index == 0 else "  "
                    canvas_obj.drawString(margin, y, f"{prefix}{line}")
                    y -= 14 if (style or "").lower() == "slim" else 15

    canvas_obj.save()


def create_job_analysis(
    db: Session,
    job_text: Optional[str],
    job_url: Optional[str],
    user_id: Optional[UUID],
) -> JobAnalysis:
    source_text = job_text or ""
    source_url = None

    if job_url:
        source_url = job_url
        source_text = fetch_job_text(job_url)

    extracted_text = normalize_text(source_text)
    ai_result = analyze_job_with_openai(extracted_text)

    ai_raw = None
    ai_model = None
    if ai_result:
        parsed = ai_result.get("parsed", {})
        keywords = parsed.get("keywords", [])
        signals = parsed.get("signals", {})
        summary = parsed.get("summary") or summarize_text(extracted_text)
        ai_raw = ai_result.get("raw")
        ai_model = ai_result.get("model")
    else:
        keywords = extract_keywords(extracted_text)
        signals = extract_signals(extracted_text)
        summary = summarize_text(extracted_text)

    analysis = JobAnalysis(
        user_id=user_id,
        source_url=source_url,
        job_text=source_text,
        extracted_text=extracted_text,
        summary=summary,
        keywords=keywords,
        signals=signals,
        ai_raw_response=ai_raw,
        ai_model=ai_model,
    )
    db.add(analysis)
    db.commit()
    db.refresh(analysis)
    return analysis


def parse_resume_text(resume_text: str) -> tuple[dict, Optional[dict], Optional[str]]:
    ai_result = parse_resume_with_openai(resume_text)
    if ai_result:
        parsed = ai_result.get("parsed", {})
        return (
            {
                "name": parsed.get("name", ""),
                "email": parsed.get("email"),
                "phone": parsed.get("phone"),
                "skills": parsed.get("skills", []),
                "summary": parsed.get("summary", ""),
                "experience": parsed.get("experience", []),
                "education": parsed.get("education", []),
            },
            ai_result.get("raw"),
            ai_result.get("model"),
        )

    lines = [line.strip() for line in resume_text.splitlines() if line.strip()]
    name = lines[0] if lines else ""
    email_match = EMAIL_PATTERN.search(resume_text)
    phone_match = PHONE_PATTERN.search(resume_text)

    skills = []
    skills_block = ""
    for index, line in enumerate(lines):
        if line.lower().startswith("skills"):
            skills_block = " ".join(lines[index + 1 : index + 4])
            break
    if skills_block:
        skills = [item.strip() for item in skills_block.split(",") if item.strip()]
    if not skills:
        skills = extract_keywords(resume_text, limit=8)

    experience = _extract_experience_section(lines)
    education = _extract_education_section(lines)

    return (
        {
            "name": name,
            "email": email_match.group(0) if email_match else None,
            "phone": phone_match.group(0) if phone_match else None,
            "skills": skills,
            "summary": "",
            "experience": experience,
            "education": education,
        },
        None,
        None,
    )


def create_resume_profile(
    db: Session,
    resume_text: str,
    file_name: Optional[str],
    user_id: Optional[UUID],
) -> ResumeProfile:
    parsed, ai_raw, ai_model = parse_resume_text(resume_text)
    profile = ResumeProfile(
        user_id=user_id,
        file_name=file_name,
        raw_text=resume_text,
        parsed_data=parsed,
        ai_raw_response=ai_raw,
        ai_model=ai_model,
    )
    db.add(profile)
    db.commit()
    db.refresh(profile)
    return profile


def update_resume_profile_data(
    db: Session,
    resume_profile: ResumeProfile,
    statement: Optional[str],
    skills: Optional[list[str]],
    experience: Optional[list[dict]],
    education: Optional[list[dict]],
) -> ResumeProfile:
    parsed = dict(resume_profile.parsed_data or {})
    if statement is not None:
        parsed["summary"] = statement
    if skills is not None:
        parsed["skills"] = skills
    if experience is not None:
        parsed["experience"] = experience
    if education is not None:
        parsed["education"] = education

    resume_profile.parsed_data = parsed
    db.add(resume_profile)
    db.commit()
    db.refresh(resume_profile)
    return resume_profile


def _tailor_bullets(bullets: list[str], keywords: list[str]) -> list[str]:
    if not bullets:
        return []

    tailored = []
    for index, bullet in enumerate(bullets):
        keyword = keywords[index % max(len(keywords), 1)] if keywords else ""
        if keyword and keyword.lower() not in bullet.lower():
            tailored.append(f"{bullet} ({keyword})")
        else:
            tailored.append(bullet)
    return tailored


def _tailor_experience(experiences: list[dict], keywords: list[str]) -> list[dict]:
    tailored = []
    for experience in experiences:
        bullets = experience.get("bullets", [])
        tailored.append({**experience, "bullets": _tailor_bullets(bullets, keywords)})
    return tailored


def _tailor_education(education: list[dict], keywords: list[str]) -> list[dict]:
    tailored = []
    for edu in education:
        bullets = edu.get("bullets", [])
        tailored.append({**edu, "bullets": _tailor_bullets(bullets, keywords[:3])})
    return tailored


def build_tailored_output(
    job_analysis: JobAnalysis,
    resume_profile: ResumeProfile,
    target_role: Optional[str],
) -> tuple[str, list[str], str, list[dict], list[dict]]:
    role = target_role or "Role"
    skills = resume_profile.parsed_data.get("skills", [])
    keywords = job_analysis.keywords

    summary = (
        f"{resume_profile.parsed_data.get('name', 'Candidate')} is a {role} with "
        f"strengths in {', '.join(skills[:3])}. Focused on {', '.join(keywords[:3])} "
        "and proven delivery across cross-functional teams."
    )

    bullets = [
        f"Aligned {role} experience with {keyword} outcomes for high-impact projects"
        for keyword in keywords[:3]
    ]
    if skills:
        bullets.append(
            f"Leveraged {', '.join(skills[:2])} to deliver measurable product results"
        )

    base_name = resume_profile.file_name or "Resume"
    base_name = re.sub(r"\.[^/.]+$", "", base_name)
    base_name = sanitize_filename(base_name)

    experiences = resume_profile.parsed_data.get("experience", [])
    education = resume_profile.parsed_data.get("education", [])

    tailored_experience = _tailor_experience(experiences, keywords)
    tailored_education = _tailor_education(education, keywords)

    return summary, bullets, base_name, tailored_experience, tailored_education


def create_tailored_resume(
    db: Session,
    job_analysis: JobAnalysis,
    resume_profile: ResumeProfile,
    target_role: Optional[str],
    user_id: Optional[UUID],
    style: Optional[str] = None,
) -> TailoredResume:
    selected_style = style or "Modern"
    summary, bullets, base_name, tailored_experience, tailored_education = build_tailored_output(
        job_analysis, resume_profile, target_role
    )

    tailored_id = uuid4()
    folder = ensure_storage_dir("tailored", str(tailored_id))
    pdf_name = f"{base_name}-tailored.pdf"
    docx_name = f"{base_name}-tailored.docx"
    pdf_path = os.path.join(folder, pdf_name)
    docx_path = os.path.join(folder, docx_name)
    relative_pdf = f"tailored/{tailored_id}/{pdf_name}"
    relative_docx = f"tailored/{tailored_id}/{docx_name}"

    name = resume_profile.parsed_data.get("name") or "Candidate"
    role = target_role or "Role"
    render_pdf(
        pdf_path,
        name,
        role,
        summary,
        bullets,
        tailored_experience,
        tailored_education,
        selected_style,
    )
    render_docx(
        docx_path,
        name,
        role,
        summary,
        bullets,
        tailored_experience,
        tailored_education,
        selected_style,
    )

    tailored = TailoredResume(
        id=tailored_id,
        user_id=user_id,
        job_analysis_id=job_analysis.id,
        resume_profile_id=resume_profile.id,
        target_role=target_role,
        style=selected_style,
        tailored_summary=summary,
        tailored_bullets=bullets,
        tailored_experience=tailored_experience,
        tailored_education=tailored_education,
        output_files=[relative_pdf, relative_docx],
    )
    db.add(tailored)
    db.commit()
    db.refresh(tailored)
    return tailored


def get_job_analysis(db: Session, analysis_id: UUID) -> Optional[JobAnalysis]:
    return db.query(JobAnalysis).filter(JobAnalysis.id == analysis_id).first()


def get_resume_profile(db: Session, profile_id: UUID) -> Optional[ResumeProfile]:
    return db.query(ResumeProfile).filter(ResumeProfile.id == profile_id).first()


def get_tailored_resume(db: Session, tailored_id: UUID) -> Optional[TailoredResume]:
    return db.query(TailoredResume).filter(TailoredResume.id == tailored_id).first()
